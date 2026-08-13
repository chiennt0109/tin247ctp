import hashlib
import io
import os
import re
import zipfile
from pathlib import Path
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from decimal import Decimal

from django.core.exceptions import PermissionDenied, ValidationError
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from assessment.models import ExamAccessGrant, ExamAttempt
from assessment.services.protected_payload import decrypt_json

MAX_VARIANTS = 8
VARIANT_CHOICES = {1, 4, 8}
FULL_PACKAGE = "exam_answers"
class ExportValidationError(ValidationError):
    pass


@dataclass(frozen=True)
class DownloadPermission:
    allowed: bool
    grant_id: int | None = None


@dataclass
class ExportContext:
    owner: object
    user_id: int
    session: object
    generated_exam: object
    blueprint: object
    blueprint_version: object
    scoring_version: object
    questions: list
    package_id: str
    attempt_id: str | None
    grant_id: int | None
    package: str
    variants: int

    @property
    def exam_code(self):
        return _safe_name(self.generated_exam.code or str(self.generated_exam.pk))

    @property
    def root(self):
        return f"GOI_DE_{self.exam_code}"


def user_download_permission(user, session):
    """Return True only for an active direct/group grant that explicitly allows downloads."""
    grants = ExamAccessGrant.objects.filter(
        session=session, is_active=True, allow_download=True,
    )
    direct = grants.filter(user=user).first()
    if direct:
        return DownloadPermission(True, direct.pk)
    group = grants.filter(group__in=user.groups.all(), user__isnull=True).order_by("pk").first()
    if group:
        return DownloadPermission(True, group.pk)
    return DownloadPermission(False)


def build_attempt_download_zip(*, attempt, user, package, variants=1):
    """Build a standardized ZIP from the existing generated-exam snapshot only."""
    if attempt.user_id != user.pk:
        raise PermissionDenied("Bạn không có quyền tải bài làm này.")
    permission = user_download_permission(user, attempt.session)
    if not permission.allowed:
        raise PermissionDenied("Tài khoản chưa được cấp quyền tải đề.")
    variants = _clean_variants(variants)
    _validate_package(package)
    attempt = _hydrate_attempt(attempt.pk)
    ctx = _context_from_attempt(attempt, package=package, variants=variants, grant_id=permission.grant_id)
    return _build_standard_zip(ctx)


def build_resource_package_zip(*, resource_package, user, package, variants=1):
    if resource_package.user_id != user.pk:
        raise PermissionDenied("Bạn không có quyền tải gói này.")
    permission = user_download_permission(user, resource_package.session)
    if not permission.allowed:
        raise PermissionDenied("Tài khoản chưa được cấp quyền tải đề.")
    variants = _clean_variants(variants)
    _validate_package(package)
    ctx = _context_from_resource_package(
        resource_package, package=package, variants=variants, grant_id=permission.grant_id,
    )
    return _build_standard_zip(ctx)


def _clean_variants(variants):
    try:
        variants = int(variants)
    except (TypeError, ValueError) as exc:
        raise ValidationError("Số mã đề không hợp lệ.") from exc
    if variants not in VARIANT_CHOICES or variants > MAX_VARIANTS:
        raise ValidationError("Chỉ được tạo 1, 4 hoặc 8 mã đề.")
    return variants


def _validate_package(package):
    if package not in {"exam", "exam_answers", "blueprint"}:
        raise ValidationError("Gói tải xuống không hợp lệ.")


def _hydrate_attempt(pk):
    return (
        ExamAttempt.objects.select_related(
            "session", "generated_exam", "blueprint", "blueprint_version",
            "generated_exam__scoring_version", "generated_exam__scoring_version__scheme",
            "generated_exam__blueprint_version", "generated_exam__blueprint_version__blueprint",
        )
        .prefetch_related(
            "generated_exam__questions__bank_question__curriculum",
            "generated_exam__questions__bank_question__outcome",
            "generated_exam__questions__bank_revision",
            "generated_exam__questions__blueprint_slot__curriculum",
            "generated_exam__questions__blueprint_slot__outcome",
            "generated_exam__blueprint_version__sections__slots__curriculum",
            "generated_exam__blueprint_version__sections__slots__outcome",
            "generated_exam__scoring_version__rules",
        )
        .get(pk=pk)
    )


def _context_from_attempt(attempt, *, package, variants, grant_id):
    exam = attempt.generated_exam
    return ExportContext(
        owner=attempt, user_id=attempt.user_id, session=attempt.session,
        generated_exam=exam, blueprint=attempt.blueprint,
        blueprint_version=exam.blueprint_version, scoring_version=exam.scoring_version,
        questions=list(exam.questions.all().order_by("order")),
        package_id=str(attempt.pk), attempt_id=str(attempt.pk), grant_id=grant_id,
        package=package, variants=variants,
    )


def _context_from_resource_package(resource_package, *, package, variants, grant_id):
    exam = resource_package.generated_exam
    return ExportContext(
        owner=resource_package, user_id=resource_package.user_id, session=resource_package.session,
        generated_exam=exam, blueprint=resource_package.blueprint,
        blueprint_version=resource_package.blueprint_version, scoring_version=exam.scoring_version,
        questions=list(exam.questions.select_related(
            "bank_question__curriculum", "bank_question__outcome", "bank_revision",
            "blueprint_slot__curriculum", "blueprint_slot__outcome",
        ).order_by("order")),
        package_id=str(resource_package.pk), attempt_id=None, grant_id=grant_id,
        package=package, variants=variants,
    )


def _build_standard_zip(ctx):
    errors = _semantic_validation_errors(ctx)
    if errors:
        raise ExportValidationError(errors)
    files = _build_export_files(ctx)
    archive_buffer = io.BytesIO()
    with zipfile.ZipFile(archive_buffer, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        for name, payload in files.items():
            archive.writestr(f"{ctx.root}/{name}", payload)
    archive_buffer.seek(0)
    return archive_buffer.getvalue()


def _build_export_files(ctx):
    code = ctx.exam_code
    include_exam = ctx.package in {"exam", FULL_PACKAGE}
    include_answers = ctx.package == FULL_PACKAGE
    include_blueprint = ctx.package in {"blueprint", FULL_PACKAGE}
    files = {}
    if include_exam:
        exam_lines = _exam_lines(ctx, include_answers=False)
        exam_docx = _form_docx_bytes(f"Đề thi {code}", exam_lines)
        files[f"01_DE_THI_{code}.docx"] = exam_docx
        files[f"01_DE_THI_{code}.pdf"] = _render_pdf(exam_docx, f"01_DE_THI_{code}.docx")
    if include_answers:
        answer_lines = _answer_lines(ctx)
        answer_docx = _form_docx_bytes(f"Đáp án {code}", answer_lines)
        files[f"02_DAP_AN_{code}.docx"] = answer_docx
        files[f"02_DAP_AN_{code}.pdf"] = _render_pdf(answer_docx, f"02_DAP_AN_{code}.docx")
    if include_blueprint:
        matrix_rows = _matrix_rows(ctx)
        spec_rows = _spec_rows(ctx, public=False)
        matrix_xlsx = _form_xlsx_bytes(ctx, matrix_rows, spec_rows)
        files[f"03_MA_TRAN_{code}.xlsx"] = matrix_xlsx
        files[f"03_MA_TRAN_{code}.pdf"] = _render_pdf(matrix_xlsx, f"03_MA_TRAN_{code}.xlsx")
        spec_docx = _form_docx_bytes(f"Bản đặc tả {code}", _rows_to_lines(spec_rows))
        files[f"04_BAN_DAC_TA_{code}.docx"] = spec_docx
        files[f"04_BAN_DAC_TA_{code}.pdf"] = _render_pdf(spec_docx, f"04_BAN_DAC_TA_{code}.docx")
    snapshot_sheets = _snapshot_sheets(ctx)
    files[f"INTERNAL/05_SNAPSHOT_{code}.xlsx"] = _xlsx_bytes(snapshot_sheets)
    validation = _validation_lines(ctx, files)
    files[f"INTERNAL/07_VALIDATION_REPORT_{code}.txt"] = "\n".join(validation).encode("utf-8")
    manifest = _manifest_lines(ctx, files)
    files[f"INTERNAL/06_MANIFEST_{code}.txt"] = "\n".join(manifest).encode("utf-8")
    files["INTERNAL/README.txt"] = "\n".join(_readme_lines(ctx)).encode("utf-8")
    return files


def _exam_lines(ctx, *, include_answers):
    lines = [
        "ĐƠN VỊ: THPT chuyên Trần Phú",
        f"KỲ THI: {ctx.session.name}",
        "MÔN: Tin học",
        f"NĂM HỌC: {getattr(ctx.session, 'opens_at', None).year if getattr(ctx.session, 'opens_at', None) else 'NEEDS_REVIEW'}",
        f"MÃ ĐỀ: {ctx.generated_exam.code}",
        f"ĐỊNH HƯỚNG: {ctx.blueprint.exam_type}",
        f"THỜI GIAN LÀM BÀI: {ctx.session.duration_minutes} phút",
        f"TỔNG ĐIỂM: {ctx.generated_exam.total_score}",
        "HƯỚNG DẪN: Thí sinh làm bài theo từng phần. Không ghi đáp án vào đề.",
        "Trang 1/1",
        "",
        "PHẦN I — TRẮC NGHIỆM NHIỀU LỰA CHỌN",
    ]
    for question in _mcq_questions(ctx):
        lines.extend(_question_exam_lines(question, include_answers=include_answers))
    lines.extend(["", "PHẦN II — TRẮC NGHIỆM ĐÚNG/SAI"])
    for question in _tf_questions(ctx):
        lines.extend(_question_exam_lines(question, include_answers=include_answers))
    manual = _manual_questions(ctx)
    if manual:
        lines.extend(["", "PHẦN III — TỰ LUẬN/THỰC HÀNH"])
        for question in manual:
            lines.extend(_question_exam_lines(question, include_answers=include_answers))
    return lines


def _question_exam_lines(question, *, include_answers):
    lines = [f"Câu {question.order}. {question.stem_snapshot}"]
    if question.options_snapshot:
        options = _ordered_options(question)
        for index, option in enumerate(options):
            lines.append(f"  {chr(65 + index)}. {_text(option)}")
        if include_answers:
            lines.append(f"  Đáp án: {_mcq_answer_label(question)}")
    elif question.statements_snapshot:
        for index, statement in enumerate(_ordered_statements(question)):
            lines.append(f"  {chr(97 + index)}) {_text(statement)}")
        if include_answers:
            answer = _tf_answer_map(question)
            lines.append("  Đáp án: " + "; ".join(f"{k}: {v}" for k, v in answer.items()))
    lines.append("")
    return lines


def _answer_lines(ctx):
    lines = [f"ĐÁP ÁN — {ctx.session.name}", f"Mã đề: {ctx.generated_exam.code}", ""]
    mcq = _mcq_questions(ctx)
    if mcq:
        lines.append("PHẦN I — TRẮC NGHIỆM NHIỀU LỰA CHỌN")
        lines.append("Câu | " + " | ".join(str(q.order) for q in mcq))
        lines.append("Đáp án | " + " | ".join(str(_mcq_answer_label(q)) for q in mcq))
        lines.append("")
    tf = _tf_questions(ctx)
    if tf:
        lines.append("PHẦN II — TRẮC NGHIỆM ĐÚNG/SAI")
        lines.append("Câu | a | b | c | d")
        for question in tf:
            answer = _tf_answer_map(question)
            lines.append(f"{question.order} | {answer.get('a', 'NEEDS_REVIEW')} | {answer.get('b', 'NEEDS_REVIEW')} | {answer.get('c', 'NEEDS_REVIEW')} | {answer.get('d', 'NEEDS_REVIEW')}")
    manual = _manual_questions(ctx)
    if manual:
        lines.extend(["", "PHẦN III — HƯỚNG DẪN CHẤM TỰ LUẬN/THỰC HÀNH"])
        for question in manual:
            payload = decrypt_json(question.protected_answer_snapshot)
            guide = payload.get("answer_guide") or payload.get("answer_key") or "NEEDS_REVIEW"
            lines.extend([f"Câu {question.order} ({question.score} điểm)", str(guide), ""])
    return lines


def _mcq_questions(ctx):
    return [q for q in ctx.questions if q.options_snapshot]


def _tf_questions(ctx):
    return [q for q in ctx.questions if q.statements_snapshot]


def _manual_questions(ctx):
    return [q for q in ctx.questions if q.bank_question.question_type in {"ESSAY", "PRACTICAL"}]


def _ordered_options(question):
    # GeneratedExamQuestion stores the already-displayed snapshot. The order
    # vector is retained solely to map answers back to canonical bank indexes.
    return list(question.options_snapshot or [])


def _ordered_statements(question):
    return list(question.statements_snapshot or [])


def _mcq_answer_label(question):
    expected = str(decrypt_json(question.protected_answer_snapshot).get("answer_key", "")).strip().upper()
    labels = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    accepted_originals = {expected}
    if expected in labels:
        accepted_originals.add(str(labels.index(expected)))
    for display_index, original_index in enumerate(question.option_order or range(len(question.options_snapshot or []))):
        if str(original_index).upper() in accepted_originals:
            return labels[display_index]
    return expected or "NEEDS_REVIEW"


def _tf_answer_map(question):
    raw = decrypt_json(question.protected_answer_snapshot).get("answer_key", {})
    if isinstance(raw, dict):
        values = [raw.get(str(i), raw.get(chr(97 + i))) for i in range(4)]
    elif isinstance(raw, list):
        values = raw[:4]
    else:
        values = []
    order = question.statement_order or list(range(len(values)))
    result = {}
    for display_index in range(4):
        try:
            value = values[int(order[display_index])]
        except (IndexError, TypeError, ValueError):
            value = None
        result[chr(97 + display_index)] = _true_false_label(value)
    return result


def _true_false_label(value):
    if isinstance(value, str):
        value = value.strip().lower()
        if value in {"true", "t", "đ", "d", "1", "yes"}:
            return "Đ"
        if value in {"false", "f", "s", "0", "no"}:
            return "S"
    if value is True:
        return "Đ"
    if value is False:
        return "S"
    return "NEEDS_REVIEW"


def _text(value):
    if isinstance(value, dict):
        return str(value.get("text") or value.get("content") or value)
    return str(value)


def _matrix_rows(ctx):
    header = [
        "TT", "Chương/Chủ đề", "Nội dung/Đơn vị kiến thức",
        "TNKQ-NLC: Biết", "TNKQ-NLC: Hiểu", "TNKQ-NLC: Vận dụng",
        "TNKQ-ĐS: Biết", "TNKQ-ĐS: Hiểu", "TNKQ-ĐS: Vận dụng",
        "Tổng: Biết", "Tổng: Hiểu", "Tổng: Vận dụng", "Tỉ lệ % điểm",
    ]
    buckets = {}
    for question in ctx.questions:
        slot = question.blueprint_slot
        curriculum = slot.curriculum or question.bank_question.curriculum
        outcome = slot.outcome or question.bank_question.outcome
        key = (
            getattr(curriculum, "source_id", "NEEDS_REVIEW"),
            getattr(outcome, "source_id", "NEEDS_REVIEW"),
        )
        bucket = buckets.setdefault(key, {
            "topic": getattr(curriculum, "topic_name", "NEEDS_REVIEW"),
            "unit": getattr(outcome, "text", "NEEDS_REVIEW"),
            "cells": {name: [] for name in header[3:12]},
            "score": Decimal("0"),
        })
        if question.statements_snapshot:
            for position, statement in zip(_positions(question), _ordered_statements(question), strict=True):
                level = _level_name(statement.get("cognitive_level") if isinstance(statement, dict) else "")
                bucket["cells"].setdefault(f"TNKQ-ĐS: {level}", []).append(position)
                bucket["cells"].setdefault(f"Tổng: {level}", []).append(position)
        elif question.bank_question.question_type not in {"ESSAY", "PRACTICAL"}:
            level = _level_name(slot.cognitive_level or question.bank_question.cognitive_level)
            positions = _positions(question)
            bucket["cells"].setdefault(f"TNKQ-NLC: {level}", []).extend(positions)
            bucket["cells"].setdefault(f"Tổng: {level}", []).extend(positions)
        bucket["score"] += Decimal(question.score)
    rows = [header]
    total_score = sum((Decimal(q.score) for q in ctx.questions), Decimal("0")) or Decimal("1")
    totals = {name: 0 for name in header[3:12]}
    for index, bucket in enumerate(buckets.values(), start=1):
        row = [index, bucket["topic"], bucket["unit"]]
        for name in header[3:12]:
            positions = bucket["cells"].get(name, [])
            totals[name] += len(positions)
            row.append(_count_cell(positions))
        row.append(f"{(bucket['score'] / total_score * 100):.2f}%")
        rows.append(row)
    total_row = ["Tổng", len(ctx.questions), f"{_command_count(ctx)} lệnh hỏi / {ctx.generated_exam.total_score} điểm"]
    for name in header[3:12]:
        total_row.append(str(totals[name]))
    total_row.append("100%")
    rows.append(total_row)
    return rows


def _positions(question):
    if question.statements_snapshot:
        return [f"Câu {question.order}{chr(97 + i)}" for i, _ in enumerate(_ordered_statements(question))]
    return [f"Câu {question.order}"]


def _count_cell(positions):
    return "" if not positions else f"{len(positions)}\n({', '.join(positions)})"


def _level_name(value):
    value = (value or "").upper()
    if "UNDER" in value or "HIEU" in value or "HIỂU" in value:
        return "Hiểu"
    if "APPL" in value or "VD" in value or "VẬN" in value:
        return "Vận dụng"
    return "Biết"


def _command_count(ctx):
    return sum(len(q.statements_snapshot) if q.statements_snapshot else 1 for q in ctx.questions)


def _spec_rows(ctx, *, public):
    rows = [[
        "TT", "Chủ đề", "Nội dung/Đơn vị kiến thức", "Curriculum_ID", "Outcome_ID",
        "Yêu cầu cần đạt", "Loại câu", "Mức tư duy", "Năng lực", "Số lệnh hỏi", "Điểm",
        "Vị trí câu/ý", "Mô tả yêu cầu đánh giá", "Question_ID được chọn",
    ]]
    for index, question in enumerate(ctx.questions, start=1):
        slot = question.blueprint_slot
        curriculum = slot.curriculum or question.bank_question.curriculum
        outcome = slot.outcome or question.bank_question.outcome
        rows.append([
            index,
            getattr(curriculum, "topic_name", "NEEDS_REVIEW"),
            getattr(outcome, "text", "NEEDS_REVIEW"),
            getattr(curriculum, "source_id", "NEEDS_REVIEW"),
            getattr(outcome, "source_id", "NEEDS_REVIEW"),
            getattr(outcome, "text", "NEEDS_REVIEW"),
            ({"ESSAY": "Tự luận", "PRACTICAL": "Thực hành"}.get(
                question.bank_question.question_type,
                "Đúng/Sai" if question.statements_snapshot else "Nhiều lựa chọn",
            )),
            slot.cognitive_level or question.bank_question.cognitive_level or "NEEDS_REVIEW",
            slot.competency or question.bank_question.competency or "NEEDS_REVIEW",
            len(question.statements_snapshot) if question.statements_snapshot else 1,
            str(question.score),
            ", ".join(_positions(question)),
            question.bank_question.source_metadata.get("assessment_description", "NEEDS_REVIEW"),
            "ẨN" if public else question.question_id_snapshot,
        ])
    return rows


def _snapshot_sheets(ctx):
    question_rows = [["order", "question_id", "source_version", "slot_id", "score", "content_hash"]]
    order_rows = [["order", "kind", "option_or_statement_order"]]
    source_rows = [["order", "question_id", "curriculum_id", "outcome_id", "family_id"]]
    for question in ctx.questions:
        question_rows.append([
            question.order, question.question_id_snapshot, question.source_version_snapshot,
            question.blueprint_slot_id, str(question.score), question.content_hash_snapshot,
        ])
        order_rows.append([
            question.order,
            ("STATEMENT" if question.statements_snapshot else
             "MANUAL" if question.bank_question.question_type in {"ESSAY", "PRACTICAL"} else "OPTION"),
            ",".join(map(str, (
                question.statement_order if question.statements_snapshot else question.option_order
            ) or [])),
        ])
        source_rows.append([
            question.order, question.question_id_snapshot,
            getattr(question.bank_question.curriculum, "source_id", ""),
            getattr(question.bank_question.outcome, "source_id", ""),
            question.bank_question.duplicate_family_id,
        ])
    return {
        "README": [["Package_ID", ctx.package_id], ["GeneratedExam_ID", ctx.generated_exam.pk], ["ExamAttempt_ID", ctx.attempt_id or ""], ["ExamSession", ctx.session.name], ["user", ctx.user_id], ["Blueprint_ID", ctx.blueprint.pk], ["Blueprint version", ctx.blueprint_version.version], ["Scoring version", ctx.scoring_version.version], ["exam code", ctx.generated_exam.code], ["orientation", ctx.blueprint.exam_type], ["seed", ctx.generated_exam.seed], ["question groups", len(ctx.questions)], ["commands", _command_count(ctx)], ["total score", str(ctx.generated_exam.total_score)], ["preview/official", "PREVIEW" if not ctx.generated_exam.is_locked else "OFFICIAL"]],
        "EXAM": [["field", "value"], ["session", ctx.session.name], ["duration", ctx.session.duration_minutes], ["code", ctx.generated_exam.code]],
        "EXAM_ITEMS": question_rows,
        "BLUEPRINT": [["id", ctx.blueprint.pk], ["name", ctx.blueprint.name], ["source", ctx.blueprint.source_blueprint_id]],
        "BLUEPRINT_SLOTS": _blueprint_slot_rows(ctx),
        "QUESTIONS_USED": question_rows,
        "OPTION_OR_STATEMENT_ORDER": order_rows,
        "SCORING_SNAPSHOT": _scoring_rows(ctx),
        "SOURCE_SNAPSHOT": source_rows,
        "AUDIT_SNAPSHOT": [["generated_at", ctx.generated_exam.generated_at.isoformat()], ["grant_id", ctx.grant_id or ""]],
        "FILE_MANIFEST": [["generated_in_zip", "see 06_MANIFEST txt"]],
        "VALIDATION_RESULTS": [["generated_in_zip", "see 07_VALIDATION_REPORT txt"]],
    }


def _blueprint_slot_rows(ctx):
    rows = [["section", "slot_id", "order", "question_type", "cognitive_level", "quantity", "score_per_item"]]
    for section in ctx.blueprint_version.sections.all().order_by("order", "id"):
        for slot in section.slots.all().order_by("order", "id"):
            rows.append([section.code, slot.pk, slot.order, slot.question_type, slot.cognitive_level, slot.quantity, str(slot.score_per_item)])
    return rows


def _scoring_rows(ctx):
    rows = [["question_type", "rule_code", "max_score", "configuration", "order"]]
    for rule in ctx.scoring_version.rules.all().order_by("order", "id"):
        rows.append([rule.question_type, rule.rule_code, str(rule.max_score), str(rule.configuration), rule.order])
    return rows


def _validation_lines(ctx, files):
    errors = _semantic_validation_errors(ctx)
    return ([f"FAIL | {error}" for error in errors] or ["PASS | SEMANTIC_SNAPSHOT_VALIDATION"])


def _order_is_valid(question):
    size = len(question.options_snapshot or question.statements_snapshot or [])
    order = (
        question.statement_order if question.statements_snapshot else question.option_order
    ) or list(range(size))
    return sorted([int(x) for x in order]) == list(range(size))


def _manifest_lines(ctx, files):
    lines = [
        f"Package_ID: {ctx.package_id}",
        f"GeneratedExam_ID: {ctx.generated_exam.pk}",
        f"ExamSession: {ctx.session.name}",
        f"Blueprint_ID/version: {ctx.blueprint.pk}/{ctx.blueprint_version.version}",
        f"Scoring version: {ctx.scoring_version.version}",
        f"Seed: {ctx.generated_exam.seed}",
        f"Created at: {ctx.generated_exam.generated_at.isoformat()}",
        f"Preview/official: {'PREVIEW' if not ctx.generated_exam.is_locked else 'OFFICIAL'}",
        "",
        "FILES:",
    ]
    for name, payload in sorted(files.items()):
        lines.append(f"- {name} | {len(payload)} bytes | sha256={hashlib.sha256(payload).hexdigest()}")
    lines.append("")
    lines.append("QUESTIONS:")
    for question in ctx.questions:
        order = question.statement_order if question.statements_snapshot else question.option_order
        lines.append(f"- {question.order}: {question.question_id_snapshot}@{question.source_version_snapshot} | order={order}")
    return lines


def _cognitive_distribution(questions):
    result = {"BIET": 0, "HIEU": 0, "VANDUNG": 0}
    for question in questions:
        if question.statements_snapshot:
            for statement in _ordered_statements(question):
                level = str(statement.get("cognitive_level") or "") if isinstance(statement, dict) else ""
                if level not in result:
                    continue
                result[level] += 1
        else:
            level = question.blueprint_slot.cognitive_level or question.bank_question.cognitive_level
            if level in result:
                result[level] += 1
    return result


def _semantic_validation_errors(ctx):
    errors = []
    questions = list(ctx.questions)
    slots = [
        slot for section in ctx.blueprint_version.sections.all()
        for slot in section.slots.all()
    ]
    ids = [q.question_id_snapshot for q in questions]
    families = [q.family_id_snapshot or q.bank_question.duplicate_family_id for q in questions]
    families = [value for value in families if value]
    positions = [q.blueprint_slot_no_snapshot or q.blueprint_slot.source_slot_no or q.order for q in questions]
    if len(ids) != len(set(ids)):
        errors.append("QUESTION_ID_DUPLICATE")
    if len(families) != len(set(families)):
        errors.append("FAMILY_ID_DUPLICATE")
    if len(positions) != len(set(positions)):
        errors.append("POSITION_DUPLICATE")
    if len(questions) != ctx.blueprint_version.expected_question_count:
        errors.append("QUESTION_GROUP_TOTAL_MISMATCH")
    expected_types = {}
    for slot in slots:
        expected_types[slot.question_type] = expected_types.get(slot.question_type, 0) + slot.quantity
    actual_types = {}
    for question in questions:
        qtype = question.bank_question.question_type
        actual_types[qtype] = actual_types.get(qtype, 0) + 1
    if expected_types != actual_types:
        errors.append("QUESTION_TYPE_DISTRIBUTION_MISMATCH")
    if sum(Decimal(q.score) for q in questions) != Decimal(ctx.generated_exam.total_score):
        errors.append("SCORE_TOTAL_MISMATCH")
    by_slot = {q.blueprint_slot_id: q for q in questions}
    for slot in slots:
        question = by_slot.get(slot.pk)
        if not question:
            errors.append("BLUEPRINT_SLOT_MISMATCH")
            continue
        if question.bank_question.question_type != slot.question_type:
            errors.append("BLUEPRINT_CELL_MISMATCH")
        if slot.curriculum_id and question.curriculum_id_snapshot != slot.curriculum.source_id:
            errors.append("CURRICULUM_MISMATCH")
        if slot.outcome_id and question.outcome_id_snapshot != slot.outcome.source_id:
            errors.append("OUTCOME_MISMATCH")
    for question in questions:
        if question.statements_snapshot:
            if not _order_is_valid(question):
                errors.append("STATEMENT_ORDER_INVALID")
            if any(not isinstance(s, dict) or not s.get("cognitive_level") for s in question.statements_snapshot):
                errors.append("STATEMENT_LEVEL_METADATA_MISSING")
        elif question.options_snapshot and not _order_is_valid(question):
            errors.append("OPTION_ORDER_INVALID")
    note = str(ctx.blueprint_version.source_snapshot.get("NOTE") or "")
    target = re.search(r"TOTAL_COMMANDS=(\d+)", note)
    if target and _command_count(ctx) != int(target.group(1)):
        errors.append("COMMAND_TOTAL_MISMATCH")
    cognitive_target = re.search(
        r"TOTAL_COMMAND_TARGET=BIET:(\d+),HIEU:(\d+),VANDUNG:(\d+)", note,
    )
    if cognitive_target:
        expected = dict(zip(("BIET", "HIEU", "VANDUNG"), map(int, cognitive_target.groups()), strict=True))
        if _cognitive_distribution(questions) != expected:
            errors.append("COGNITIVE_DISTRIBUTION_MISMATCH")
    return list(dict.fromkeys(errors))


def _readme_lines(ctx):
    return [
        f"Gói đề: {ctx.root}",
        f"Kỳ thi: {ctx.session.name}",
        f"Mã đề: {ctx.generated_exam.code}",
        "Các file được dựng lại từ GeneratedExam và snapshot đã lưu trong database.",
        "Không đọc Excel master, không chọn lại câu hỏi, không trừ thêm lượt khi tải lại.",
        "Không có file JSON trong gói ZIP.",
    ]


def _rows_to_lines(rows):
    return [" | ".join(map(str, row)) for row in rows]


def _xlsx_bytes(sheets):
    wb = Workbook()
    first = True
    for name, rows in sheets.items():
        ws = wb.active if first else wb.create_sheet()
        first = False
        ws.title = str(name)[:31]
        for row in rows:
            ws.append(list(row))
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="EAF2FF")
        for column_cells in ws.columns:
            max_len = max((len(str(cell.value or "")) for cell in column_cells), default=10)
            ws.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(max_len + 2, 12), 60)
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _safe_name(value):
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(value)).strip("._")
    return cleaned or "EXAM"


def _form_docx_bytes(title, lines):
    """Build the editable assessment form using the approved 07_EXPORT styling."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise ExportValidationError("MISSING_DOCX_RENDERER: install python-docx") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    normal = document.styles["Normal"]
    normal.font.name = "Carlito"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Carlito")

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(10)
    run = heading.add_run(title.upper())
    run.bold = True
    run.font.name = "Carlito"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(31, 78, 121)
    for line in lines:
        text = str(line)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(text)
        if text.startswith(("PHẦN ", "I. ", "II. ", "III. ", "IV. ")):
            run.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
            paragraph.paragraph_format.keep_with_next = True
        elif text.startswith("Câu "):
            run.bold = True
            paragraph.paragraph_format.keep_with_next = True
        elif text.startswith(("A. ", "B. ", "C. ", "D. ", "a) ", "b) ", "c) ", "d) ")):
            paragraph.paragraph_format.left_indent = Cm(0.6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Trang ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _style_form_sheet(sheet, *, title=None):
    dark_blue = "1F4E78"
    sheet.freeze_panes = "A2"
    sheet.sheet_view.showGridLines = False
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.fitToWidth = 1
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_margins.left = sheet.page_margins.right = 0.25
    sheet.page_margins.top = sheet.page_margins.bottom = 0.5
    for cell in sheet[1]:
        cell.font = Font(name="Carlito", size=11, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=dark_blue)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    sheet.row_dimensions[1].height = 32
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.font = Font(name="Carlito", size=11)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            if cell.row % 2 == 0:
                cell.fill = PatternFill("solid", fgColor="F4F8FB")
    for column_cells in sheet.columns:
        letter = get_column_letter(column_cells[0].column)
        maximum = max((len(str(cell.value or "")) for cell in column_cells), default=10)
        sheet.column_dimensions[letter].width = min(max(maximum + 2, 10), 45)
    sheet.auto_filter.ref = sheet.dimensions
    sheet.print_title_rows = "1:1"
    sheet.print_area = sheet.dimensions


def _form_xlsx_bytes(ctx, matrix_rows, spec_rows):
    workbook = Workbook()
    workbook.remove(workbook.active)
    info = workbook.create_sheet("THONG_TIN")
    info.append(["Thông tin", "Giá trị"])
    for key, value in (
        ("Mã đề", ctx.generated_exam.code),
        ("Blueprint", ctx.blueprint.source_blueprint_id or str(ctx.blueprint.pk)),
        ("Phiên bản blueprint", ctx.blueprint_version.version),
        ("Thời gian (phút)", ctx.session.duration_minutes),
        ("Tổng lệnh hỏi", _command_count(ctx)),
        ("Tổng điểm", float(ctx.generated_exam.total_score)),
        ("Random Seed", ctx.generated_exam.seed),
    ):
        info.append([key, value])
    matrix = workbook.create_sheet("MA_TRAN")
    for row in matrix_rows:
        matrix.append(row)
    spec = workbook.create_sheet("BAN_DAC_TA")
    for row in spec_rows:
        spec.append(row)
    questions = workbook.create_sheet("DANH_SACH_CAU")
    questions.append([
        "SLOT_NO", "QUESTION_ID", "VERSION", "TYPE", "COGNITIVE", "CURRICULUM_ID",
        "OUTCOME_ID", "FAMILY_ID", "STATUS", "SCORE", "ORDER", "SEED", "SHUFFLE",
        "TIME_SECONDS", "DIFFICULTY", "STEM",
    ])
    question_rows = [[
        q.order, q.question_id_snapshot, q.source_version_snapshot,
        q.bank_question.question_type, q.bank_question.cognitive_level,
        q.curriculum_id_snapshot, q.outcome_id_snapshot, q.family_id_snapshot,
        q.bank_question.process_status, "", ",".join(map(str, q.option_order or q.statement_order)),
        ctx.generated_exam.seed, q.bank_question.shuffle_allowed,
        q.bank_question.estimated_time_seconds, q.bank_question.difficulty, q.stem_snapshot,
    ] for q in ctx.questions]
    for row in question_rows:
        questions.append(row)
    calculations = workbook.create_sheet("TINH_TOAN")
    calculations.append(["Chỉ số", "Giá trị"])
    calculations.append(["Tổng nhóm câu", f"=COUNTA(DANH_SACH_CAU!B2:B{len(question_rows) + 1})"])
    calculations.append(["Tổng lệnh hỏi", _command_count(ctx)])
    calculations.append(["Tổng điểm", float(ctx.generated_exam.total_score)])
    for sheet in workbook.worksheets:
        _style_form_sheet(sheet)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _render_pdf(document_bytes, filename):
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        raise ExportValidationError("MISSING_PDF_RENDERER: LibreOffice/soffice is required")
    with tempfile.TemporaryDirectory(prefix="assessment-export-") as directory:
        workdir = Path(directory)
        source = workdir / filename
        home = workdir / "home"
        cache = workdir / "cache"
        config = workdir / "config"
        profile = workdir / "libreoffice-profile"
        for path in (home, cache, config, profile):
            path.mkdir(mode=0o700)
        source.write_bytes(document_bytes)
        # Web workers commonly run with HOME=/var/www, which is not writable.
        # LibreOffice otherwise fails before conversion while creating its user
        # profile/dconf cache. Give every conversion an isolated writable home
        # and profile; the unique profile also prevents concurrent downloads
        # from sharing LibreOffice lock files.
        environment = os.environ.copy()
        # `gen` is an X11 VCL plugin despite its generic-sounding name. On a
        # server without an X display it produces "Can't open display" even
        # with --headless. `svp` is LibreOffice's genuinely headless plugin.
        environment.pop("DISPLAY", None)
        environment.pop("WAYLAND_DISPLAY", None)
        environment.update({
            "HOME": str(home),
            "XDG_CACHE_HOME": str(cache),
            "XDG_CONFIG_HOME": str(config),
            "SAL_USE_VCLPLUGIN": "svp",
        })
        command = [
            executable,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile.as_uri()}",
            "--convert-to", "pdf",
            "--outdir", directory,
            str(source),
        ]
        try:
            result = subprocess.run(
                command, capture_output=True, text=True, timeout=120,
                check=False, env=environment, cwd=directory,
            )
        except subprocess.TimeoutExpired as exc:
            raise ExportValidationError("PDF_RENDER_TIMEOUT: LibreOffice exceeded 120 seconds") from exc
        output = source.with_suffix(".pdf")
        if result.returncode or not output.is_file() or not output.read_bytes().startswith(b"%PDF"):
            raise ExportValidationError(
                "PDF_RENDER_FAILED: " + (result.stderr or result.stdout or "unknown renderer error")
            )
        return output.read_bytes()
